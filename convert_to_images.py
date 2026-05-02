"""
Convert tables, formula code blocks, and diagrams in paper.docx to PNG images.

Strategy:
  - Tables      → styled matplotlib table images (dark header, alternating rows)
  - Source Code → syntax-highlighted code-block images (dark terminal style)
  - The replacement uses lxml in-place XML surgery so all other formatting
    (headings, body text, existing PNG figures) is preserved exactly.
"""

import io, copy, warnings
warnings.filterwarnings("ignore")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import rcParams
rcParams["font.family"] = "monospace"

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT  = "/home/ubuntu/ky/agents/qlib_paper/paper.docx"
OUTPUT = "/home/ubuntu/ky/agents/qlib_paper/paper_images.docx"

doc = Document(INPUT)

# ─── helpers ─────────────────────────────────────────────────────────────────

def render_table(rows, dpi=150):
    """Render a list-of-list table to PNG bytes with styled headers."""
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows)
    data   = [r + [""] * (n_cols - len(r)) for r in rows]

    fig_w = min(16, max(7, n_cols * 2.2))
    fig_h = max(1.0, n_rows * 0.52 + 0.3)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")

    has_header = n_rows > 1
    cell_text  = data[1:] if has_header else data
    col_labels = data[0]  if has_header else None

    tbl = ax.table(
        cellText=cell_text,
        colLabels=col_labels,
        loc="center",
        cellLoc="left",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(8.5)
    tbl.scale(1, 1.75)

    # header row styling
    if has_header:
        for j in range(n_cols):
            c = tbl[0, j]
            c.set_facecolor("#1F3864")
            c.set_text_props(color="white", fontweight="bold", fontsize=9)
            c.set_edgecolor("white")

    # alternating body rows
    for i in range(1, len(cell_text) + 1):
        for j in range(n_cols):
            c = tbl[i, j]
            c.set_facecolor("#EEF2F7" if i % 2 == 0 else "white")
            c.set_edgecolor("#CCCCCC")

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=dpi,
                facecolor="white", edgecolor="none")
    buf.seek(0)
    plt.close(fig)
    return buf.getvalue()


def render_code(lines, dpi=150):
    """Render code lines to a PNG with a dark-terminal background."""
    if not lines:
        return None

    max_chars = max(len(l) for l in lines) + 4
    fig_w = min(14, max(6, max_chars * 0.085))
    fig_h = max(0.8, len(lines) * 0.28 + 0.4)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    fig.patch.set_facecolor("#1E1E2E")
    ax.set_facecolor("#1E1E2E")
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    # top bar
    ax.add_patch(mpatches.FancyBboxPatch(
        (0, 0.88), 1, 0.12, transform=ax.transAxes,
        boxstyle="square,pad=0", facecolor="#313244", edgecolor="none",
    ))
    for x, col in [(0.03, "#FF5F57"), (0.065, "#FEBC2E"), (0.10, "#28C840")]:
        ax.add_patch(plt.Circle((x, 0.94), 0.018, color=col,
                                transform=ax.transAxes))

    total = len(lines)
    for idx, line in enumerate(lines):
        y = 0.82 - idx * (0.78 / max(total, 1))
        # line-number gutter
        ax.text(0.015, y, f"{idx+1:3d}", transform=ax.transAxes,
                fontsize=8, color="#6C7086",
                fontfamily="monospace", va="top")
        # code content — colour keywords
        display = line.replace("\t", "    ")
        kw_color = "#CBA6F7"  # purple for keywords
        str_color = "#A6E3A1"  # green for strings/values
        default_color = "#CDD6F4"
        color = default_color
        for kw in ("def ", "class ", "import ", "from ", "return ", "for ",
                   "if ", "else:", "elif ", "with ", "as ", "in "):
            if kw in display:
                color = kw_color
                break
        for ch in ('"', "'", "#"):
            if ch in display:
                color = str_color
                break
        ax.text(0.07, y, display, transform=ax.transAxes,
                fontsize=8.5, color=color,
                fontfamily="monospace", va="top",
                clip_on=True)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=dpi,
                facecolor="#1E1E2E", edgecolor="none")
    buf.seek(0)
    plt.close(fig)
    return buf.getvalue()


def make_centered_image_para(doc, image_bytes, width_inches=6.2):
    """
    Add a picture to the doc (appends to body), then detach and return
    the raw <w:p> lxml element for in-place insertion.
    """
    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_obj.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
    p_xml = p_obj._p
    p_xml.getparent().remove(p_xml)   # detach from end of body
    return p_xml


def replace_element(old_xml, new_xml):
    """Replace old_xml with new_xml in its parent."""
    parent = old_xml.getparent()
    idx    = list(parent).index(old_xml)
    parent.remove(old_xml)
    parent.insert(idx, new_xml)


# ─── 1. REPLACE TABLES ───────────────────────────────────────────────────────

print(f"Converting {len(doc.tables)} tables …")
# Snapshot table list before modifying the tree
tables = list(doc.tables)

for i, table in enumerate(tables):
    rows = [[cell.text.strip() for cell in row.cells]
            for row in table.rows]
    img_bytes = render_table(rows)
    img_para  = make_centered_image_para(doc, img_bytes, width_inches=6.2)
    replace_element(table._tbl, img_para)
    print(f"  [{i+1}/{len(tables)}] Table {i} → image "
          f"({len(rows)}r × {len(table.columns)}c)")


# ─── 2. REPLACE SOURCE CODE BLOCKS ───────────────────────────────────────────
# Collect runs of consecutive 'Source Code' paragraphs (one block = one image)
print("\nFinding Source Code blocks …")

body_paras = doc.paragraphs   # live list after table removal

# Group consecutive Source Code paragraphs by their lxml elements
groups = []   # each group: list of (para_obj, lxml_p)
current = []
for para in body_paras:
    if para.style.name == "Source Code":
        current.append(para)
    else:
        if current:
            groups.append(current)
            current = []
if current:
    groups.append(current)

print(f"  Found {len(groups)} code block(s)")

for gi, group in enumerate(groups):
    lines = []
    for para in group:
        lines.append(para.text)

    img_bytes = render_code(lines)
    if img_bytes is None:
        continue

    # Build image paragraph and insert before first para of the group
    img_para = make_centered_image_para(doc, img_bytes, width_inches=5.8)
    first_p  = group[0]._p
    parent   = first_p.getparent()
    idx      = list(parent).index(first_p)

    # Remove all paragraphs in the group
    for para in group:
        parent.remove(para._p)

    # Insert image paragraph at the original position
    parent.insert(idx, img_para)
    print(f"  Code block {gi+1}: {len(lines)} lines → image")


# ─── 3. SAVE ─────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\nSaved → {OUTPUT}")

# Quick verification
doc2 = Document(OUTPUT)
remaining_tables = len(doc2.tables)
drawings_count   = len(doc2.element.body.findall(".//" + qn("w:drawing")))
print(f"Verification: tables remaining={remaining_tables}, "
      f"drawings in body={drawings_count}")
